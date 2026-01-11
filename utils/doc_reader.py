"""
DOC文件读取模块
用于读取.doc文件并提取文本内容
"""
import subprocess
import os
import tempfile
from typing import List, Dict
import re
from html import unescape
try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False


class DocReader:
    """DOC文件读取器"""
    
    def __init__(self):
        """初始化读取器"""
        pass
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本，移除HTML标签和非法字符
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        if not text:
            return ""
        
        # 移除HTML标签
        text = re.sub(r'<[^>]+>', '', text)
        
        # HTML实体解码
        try:
            text = unescape(text)
        except:
            pass
        
        # 移除控制字符（保留换行符、制表符和常见的中文标点）
        # 保留可打印字符、换行符、制表符
        cleaned = []
        for char in text:
            code = ord(char)
            # 保留可打印字符、换行符(\n)、制表符(\t)、回车符(\r)
            if (32 <= code <= 126) or (code in [9, 10, 13]) or (code >= 160):
                cleaned.append(char)
        
        text = ''.join(cleaned)
        
        # 移除多余的空白字符（但保留换行）
        text = re.sub(r'[ \t]+', ' ', text)  # 多个空格/制表符合并为一个空格
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # 多个换行符合并为两个
        
        # 移除行首行尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join([line for line in lines if line])
        
        return text.strip()
    
    def read_doc(self, filepath: str) -> str:
        """
        读取.doc或.docx文件内容
        
        Args:
            filepath: .doc或.docx文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            # 如果是.docx文件，使用python-docx库（更可靠）
            if filepath.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(filepath)
                    content_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            content_parts.append(para.text)
                    # 也读取表格中的文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content_parts.append(cell.text)
                    content = '\n'.join(content_parts)
                    if content and len(content.strip()) > 0:
                        return self._clean_text(content)
                except ImportError:
                    raise Exception("需要安装python-docx库来处理.docx文件: pip install python-docx")
                except Exception as e:
                    raise Exception(f"读取DOCX文件失败: {str(e)}")
            
            # 对于.doc文件，尝试多种方法
            # 方法1: 尝试使用textutil转换为临时txt文件，然后读取
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as tmp_file:
                tmp_path = tmp_file.name
            
            try:
                # 使用textutil转换为txt
                result = subprocess.run(
                    ['textutil', '-convert', 'txt', '-output', tmp_path, filepath],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0:
                    # 尝试多种编码读取
                    encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'macroman']
                    for encoding in encodings:
                        try:
                            with open(tmp_path, 'r', encoding=encoding, errors='ignore') as f:
                                content = f.read()
                            # 检查是否包含可读的中文内容
                            if content and len(content.strip()) > 100:
                                # 简单检查：如果包含常见中文字符，认为解码成功
                                chinese_chars = ['的', '是', '在', '有', '和', '与', '人', '法', '案', '判', '诉', '书', '原', '被']
                                if any(char in content for char in chinese_chars):
                                    content = self._clean_text(content)
                                    return content
                        except:
                            continue
            finally:
                # 清理临时文件
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            
            # 如果上述方法都失败，尝试HTML格式
            result_html = subprocess.run(
                ['textutil', '-convert', 'html', '-stdout', filepath],
                capture_output=True,
                timeout=30
            )
            
            if result_html.returncode == 0:
                # 使用chardet检测编码
                html_content = None
                if HAS_CHARDET:
                    detected = chardet.detect(result_html.stdout)
                    if detected and detected['encoding']:
                        try:
                            html_content = result_html.stdout.decode(detected['encoding'], errors='ignore')
                        except:
                            pass
                
                # 如果chardet失败，尝试多种编码解码HTML
                if not html_content:
                    encodings = ['utf-8', 'gbk', 'gb2312', 'latin1', 'cp1252', 'macroman', 'big5']
                    for encoding in encodings:
                        try:
                            html_content = result_html.stdout.decode(encoding, errors='ignore')
                            # 检查是否包含HTML标签和中文内容
                            if html_content and ('<html' in html_content.lower() or '<body' in html_content.lower()):
                                # 检查是否包含中文字符
                                chinese_chars = ['的', '是', '在', '有', '和', '与', '人', '法', '案', '判', '诉', '书']
                                if any(char in html_content for char in chinese_chars):
                                    break
                        except:
                            continue
                
                if html_content:
                    # 使用BeautifulSoup解析HTML
                    try:
                        from bs4 import BeautifulSoup
                        # 尝试不同的解析器
                        for parser in ['html.parser', 'lxml', 'html5lib']:
                            try:
                                soup = BeautifulSoup(html_content, parser)
                                content = soup.get_text(separator='\n', strip=True)
                                # 检查是否包含可读内容
                                if content and len(content.strip()) > 100:
                                    # 清理文本
                                    content = self._clean_text(content)
                                    return content
                            except:
                                continue
                    except ImportError:
                        pass
                    
                    # 如果BeautifulSoup失败，使用正则表达式移除HTML标签
                    import re
                    content = re.sub(r'<[^>]+>', '', html_content)
                    if content and len(content.strip()) > 100:
                        content = self._clean_text(content)
                        return content
            
            # 如果HTML格式失败，尝试TXT格式
            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-stdout', filepath],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0:
                encodings = ['utf-8', 'gbk', 'gb2312', 'latin1', 'cp1252', 'macroman']
                for encoding in encodings:
                    try:
                        content = result.stdout.decode(encoding, errors='ignore')
                        if content and len(content.strip()) > 100:
                            content = self._clean_text(content)
                            return content
                    except:
                        continue
            
            raise Exception("无法读取DOC文件内容：所有解码方式都失败")
                    
        except Exception as e:
            raise Exception(f"读取DOC文件失败: {str(e)}")
    
    def parse_case_from_doc(self, filepath: str) -> Dict:
        """
        从.doc文件中解析案例信息
        
        Args:
            filepath: .doc文件路径
            
        Returns:
            案例信息字典，包含：
                - title: 案例标题（从文件名提取）
                - case_text: 案例内容
                - judge_decision: 法官判决（如果能够识别）
                - case_date: 案例日期（如果能够识别）
        """
        # 从文件名提取标题
        filename = os.path.basename(filepath)
        title = os.path.splitext(filename)[0]
        
        # 读取文档内容
        content = self.read_doc(filepath)
        
        # 尝试提取法官判决部分（包含判理和结论）
        judge_decision = ""
        case_text = content
        
        # 先查找判理部分（"本院认为"、"本院经审查认为"等）
        # 注意：有些文档使用逗号，有些使用冒号
        # 对于二审判决书，可能有多个"本院认为"，应该使用最后一个
        reasoning_patterns = [
            r'本院经审查认为[，,：:]',  # "本院经审查认为，"或"本院经审查认为："（用于再审审查）
            r'本院经审理认为[，,：:]',  # "本院经审理认为，"
            r'本院审理认为[，,：:]',  # "本院审理认为，"
            r'本院认为[，,：:]',  # "本院认为，"
        ]
        
        reasoning_text = ""
        reasoning_start = -1
        reasoning_end = -1
        
        # 找到所有"本院认为"的位置
        # 对于二审判决书，可能有多个"本院认为"
        # 策略：使用第一个"本院认为"（通常是二审的判理开始），确保案例内容不包含判决
        for pattern in reasoning_patterns:
            matches = list(re.finditer(pattern, content))
            if matches:
                # 使用第一个匹配（确保案例内容不包含判决部分）
                first_match = matches[0]
                reasoning_start = first_match.start()  # 从"本院认为"开始，而不是从其后开始
                
                # 查找从这个位置到"本判决为终审判决"或"审判长"之间的完整内容
                remaining_content = content[reasoning_start:]
                # 先查找结束标记
                end_match = re.search(r'本判决为终审判决|本裁定为终审|审判长|书记员', remaining_content)
                if end_match:
                    # 提取到结束标记之前
                    reasoning_text = remaining_content[:end_match.start()].strip()
                else:
                    # 如果没有找到结束标记，查找"判决如下"作为分界
                    next_decision_match = re.search(r'判决如下[：:]|裁定如下[：:]', remaining_content)
                    if next_decision_match:
                        reasoning_end = reasoning_start + next_decision_match.start()
                        reasoning_text = content[reasoning_start:reasoning_end].strip()
                    else:
                        # 如果都没有找到，提取到文档末尾
                        reasoning_text = remaining_content.strip()
                break
        
        # 再查找结论部分（"判决如下"、"裁定如下"）
        # 对于二审判决书，可能有多个"判决如下"
        # 如果存在多个"判决如下"，应该从第一个开始提取（包含一审判决和二审判理+结论）
        decision_patterns = [
            r'判决如下[：:]',  # "判决如下："
            r'裁定如下[：:]',  # "裁定如下："（用于裁定书）
            r'裁判结果[：:]',  # "裁判结果："
        ]
        
        decision_text = ""
        decision_start = -1
        
        # 找到所有"判决如下"的位置
        all_decision_matches = []
        for pattern in decision_patterns:
            matches = list(re.finditer(pattern, content))
            all_decision_matches.extend(matches)
        
        if all_decision_matches:
            # 如果有多个"判决如下"，使用第一个（包含完整的判决内容）
            # 如果只有一个，也使用它
            first_match = min(all_decision_matches, key=lambda m: m.start())
            decision_start = first_match.end()
            
            # 提取从第一个"判决如下"到"本判决为终审判决"或"本裁定为终审"之间的内容
            # 注意：不能只匹配"本判决"，因为判决内容中可能包含"在本判决生效之日起"等表述
            remaining_content = content[decision_start:]
            # 优先匹配"本判决为终审判决"或"本裁定为终审"
            end_match = re.search(r'本判决为终审判决|本裁定为终审|本判决为终审|本裁定为终审裁定', remaining_content)
            if end_match:
                decision_text = remaining_content[:end_match.start()].strip()
            else:
                # 如果没有找到明确的结束标记，尝试匹配"审判长"或"书记员"（通常在判决书末尾）
                end_match2 = re.search(r'审判长|书记员', remaining_content)
                if end_match2:
                    decision_text = remaining_content[:end_match2.start()].strip()
                else:
                    # 如果都没有找到，提取到文档末尾
                    decision_text = remaining_content.strip()
        
        # 合并判理和结论
        # 确定案例内容的结束位置：应该是第一个判决标记（"本院认为"或"判决如下"/"裁定如下"）之前
        # 对于有多个"判决如下"的情况，应该使用最后一个"判决如下"之前的内容
        
        # 找到所有可能的判决开始位置
        all_judge_starts = []
        if reasoning_start >= 0:
            all_judge_starts.append(('reasoning', reasoning_start))
        if decision_start >= 0:
            all_judge_starts.append(('decision', decision_start))
        
        # 找到最早的判决开始位置作为案例内容的结束位置
        case_text_end_pos = -1
        if all_judge_starts:
            case_text_end_pos = min([pos for _, pos in all_judge_starts])
        
        # 合并判理和结论
        if reasoning_text and decision_text:
            if reasoning_start < decision_start:
                # 判理在结论之前，合并两者
                judge_decision = reasoning_text + "\n\n" + decision_text
            else:
                # 结论在判理之前（较少见），只使用结论
                judge_decision = decision_text
        elif reasoning_text:
            # 只有判理，没有明确的结论
            judge_decision = reasoning_text
        elif decision_text:
            # 只有结论，没有判理
            judge_decision = decision_text
        else:
            judge_decision = ""
        
        # 设置案例内容：应该是第一个判决标记之前的所有内容
        if case_text_end_pos >= 0:
            case_text = content[:case_text_end_pos].strip()
        else:
            case_text = content
        
        # 额外检查：如果案例内容中仍包含判决关键词，进一步清理
        # 找到案例内容中最后一个判决关键词的位置
        if case_text:
            last_benyuan = case_text.rfind('本院认为')
            last_caiding = case_text.rfind('裁定如下')
            last_panjue = case_text.rfind('判决如下')
            last_judge_keyword_pos = max(last_benyuan, last_caiding, last_panjue)
            
            if last_judge_keyword_pos >= 0:
                # 如果案例内容中仍有判决关键词，截取到最后一个关键词之前
                case_text = case_text[:last_judge_keyword_pos].strip()
        
        # 尝试提取日期（优先提取文档末尾的日期，通常是判决日期）
        case_date = ""
        
        # 方法1: 尝试匹配中文数字日期（如：二〇二五年十月三十一日）
        chinese_date_pattern = r'[二三四五六七八九〇零一二三四五六七八九十]+年[一二三四五六七八九十]+月[一二三四五六七八九十]+[日号]'
        chinese_date_match = re.search(chinese_date_pattern, content)
        if chinese_date_match:
            chinese_date = chinese_date_match.group(0)
            # 转换中文数字为阿拉伯数字
            chinese_to_arabic = {
                '〇': '0', '零': '0', '一': '1', '二': '2', '三': '3', '四': '4',
                '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'
            }
            
            # 提取年份
            year_match = re.search(r'([二三四五六七八九〇零一二三四五六七八九十]+)年', chinese_date)
            if year_match:
                year_chinese = year_match.group(1)
                year = ''.join([chinese_to_arabic.get(c, c) for c in year_chinese])
                if len(year) == 4:
                    # 提取月份
                    month_match = re.search(r'年([一二三四五六七八九十]+)月', chinese_date)
                    if month_match:
                        month_chinese = month_match.group(1)
                        if month_chinese == '十':
                            month = '10'
                        elif month_chinese.startswith('十'):
                            month = '1' + ''.join([chinese_to_arabic.get(c, '') for c in month_chinese[1:]])
                        elif month_chinese == '十一':
                            month = '11'
                        elif month_chinese == '十二':
                            month = '12'
                        else:
                            month = chinese_to_arabic.get(month_chinese, '1')
                        
                        # 提取日期
                        day_match = re.search(r'月([一二三四五六七八九十]+)[日号]', chinese_date)
                        if day_match:
                            day_chinese = day_match.group(1)
                            if day_chinese == '十':
                                day = '10'
                            elif day_chinese.startswith('十') and len(day_chinese) > 1:
                                if day_chinese == '十一':
                                    day = '11'
                                elif day_chinese == '十二':
                                    day = '12'
                                elif day_chinese == '十三':
                                    day = '13'
                                elif day_chinese == '十四':
                                    day = '14'
                                elif day_chinese == '十五':
                                    day = '15'
                                elif day_chinese == '十六':
                                    day = '16'
                                elif day_chinese == '十七':
                                    day = '17'
                                elif day_chinese == '十八':
                                    day = '18'
                                elif day_chinese == '十九':
                                    day = '19'
                                else:
                                    day = '1' + chinese_to_arabic.get(day_chinese[-1], '0')
                            elif day_chinese.startswith('二') and len(day_chinese) > 1:
                                if day_chinese == '二十':
                                    day = '20'
                                elif day_chinese == '二十一':
                                    day = '21'
                                elif day_chinese == '二十二':
                                    day = '22'
                                elif day_chinese == '二十三':
                                    day = '23'
                                elif day_chinese == '二十四':
                                    day = '24'
                                elif day_chinese == '二十五':
                                    day = '25'
                                elif day_chinese == '二十六':
                                    day = '26'
                                elif day_chinese == '二十七':
                                    day = '27'
                                elif day_chinese == '二十八':
                                    day = '28'
                                elif day_chinese == '二十九':
                                    day = '29'
                                else:
                                    day = '2' + chinese_to_arabic.get(day_chinese[-1], '0')
                            elif day_chinese.startswith('三') and len(day_chinese) > 1:
                                if day_chinese == '三十':
                                    day = '30'
                                elif day_chinese == '三十一':
                                    day = '31'
                                else:
                                    day = '3' + chinese_to_arabic.get(day_chinese[-1], '0')
                            else:
                                day = chinese_to_arabic.get(day_chinese, '1')
                            
                            case_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        
        # 方法2: 如果中文日期未匹配，尝试匹配阿拉伯数字日期
        if not case_date:
            # 优先从文档末尾查找日期（通常是判决日期）
            date_pattern = r'(\d{4})[年\-/](\d{1,2})[月\-/](\d{1,2})[日]?'
            # 从后往前搜索，找到最后一个匹配的日期
            matches = list(re.finditer(date_pattern, content))
            if matches:
                # 使用最后一个匹配（通常是判决日期）
                date_match = matches[-1]
                year, month, day = date_match.groups()
                case_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        
        return {
            'title': title,
            'case_text': case_text.strip(),
            'judge_decision': judge_decision.strip() if judge_decision else '',
            'case_date': case_date
        }
    
    def read_multiple_docs(self, directory: str, limit: int = None) -> List[Dict]:
        """
        批量读取目录下的所有.doc文件
        
        Args:
            directory: 包含.doc文件的目录
            limit: 限制读取的文件数量（用于测试）
            
        Returns:
            案例信息列表
        """
        cases = []
        doc_files = [f for f in os.listdir(directory) if f.endswith('.doc')]
        
        if limit:
            doc_files = doc_files[:limit]
        
        for filename in doc_files:
            filepath = os.path.join(directory, filename)
            try:
                case = self.parse_case_from_doc(filepath)
                cases.append(case)
                print(f"✓ 已读取: {filename}")
            except Exception as e:
                print(f"✗ 读取失败 {filename}: {str(e)}")
        
        return cases

